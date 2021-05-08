# -*- coding: utf-8 -*-


from docx import Document
from docx.shared import Inches

#制作docx
#data: dict_list
def make_docx(data, fileAllPath):
    document = Document()

    document.add_heading('文档总标题', 1)

    # 一级标题
    document.add_heading('一、表格', level=1)

    # 添加表格，增加style（如网格）
    #先添加一行标题行，后面再一行一行增加

    rows_dict = {
        "time": "时间",
        "srcIP": "源IP",
        "sPort": "源端口",
    }

    table = document.add_table(rows=1, cols=len(rows_dict), style='Table Grid')

    # 写标题行的内容
    hdr_cells = table.rows[0].cells

    for i, rows in enumerate(rows_dict.values()):
        hdr_cells[i].text = rows

    # 取出前十条数据一行一行放到表格里
    for data_detail in data[0:9]:
        row_cells = table.add_row().cells
        i = 0
        for key, value in rows_dict.items():
            row_cells[i].text = str(data_detail.get(key))
            i += 1

    # 统计不同源ip的日志个数
    srcIP_list = []
    srcIP_dict = {}

    for data_detail in data:
        srcIP = data_detail.get("srcIP")
        if srcIP not in srcIP_list:
            srcIP_list.append(srcIP)
            srcIP_dict[srcIP] = 0
        else:
            srcIP_dict[srcIP] += 1

    #document.add_heading('二、不同源ip的日志柱状图', level=1)

    # 图片插入到文档里,需要该目录下有图片
    #document.add_picture('/tmp/pieChart.jpg', width=Inches(5))

    document.save(fileAllPath)